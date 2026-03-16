"""
╔══════════════════════════════════════════════════════════════════════╗
║  MASHRAB MOLIYA v21.3 — NAKLADNOY GENERATOR                        ║
║  Mashrab formatiga AYNAN mos: MIJOZ MA'LUMOTLARI + tovar jadvali   ║
║  ✅ Word (.docx) — tahrirlash uchun                                ║
║  ✅ Excel (.xlsx) — buxgalteriya uchun                             ║
║  ✅ PDF — chop etish uchun                                         ║
║  ✅ Imzo + Muhr + INN + Manzil + Telefon                           ║
╚══════════════════════════════════════════════════════════════════════╝
"""
from __future__ import annotations
import io, logging
from datetime import datetime
from decimal import Decimal
from typing import Any
import pytz

log = logging.getLogger(__name__)
TZ  = pytz.timezone("Asia/Tashkent")


def _hozir() -> str:
    return datetime.now(TZ).strftime("%d.%m.%Y")

def _hozir_full() -> str:
    return datetime.now(TZ).strftime("%d.%m.%Y %H:%M")

def _pul(v: Any) -> str:
    try:
        d = Decimal(str(v).replace(",","").strip()) if not isinstance(v, Decimal) else v
        return f"{d:,.0f}"
    except Exception: return "0"

def nakladnoy_nomeri() -> str:
    return datetime.now(TZ).strftime("%Y%m%d%H%M%S")


# ════════════ WORD — SIZNING FORMATGA AYNAN MOS ════════════

def yaratish_word(data: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc     = Document()
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin  = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    KOK = RGBColor(0x1a, 0x56, 0xdb)
    QORA = RGBColor(0x00, 0x00, 0x00)
    KULRANG = RGBColor(0x6b, 0x72, 0x80)
    QIZIL = RGBColor(0xDC, 0x26, 0x26)

    inv_no   = data.get("invoice_number", "N/A")
    sana     = data.get("date", _hozir())
    dokon    = data.get("dokon_nomi", "Mashrab Moliya")
    dokon_tel = data.get("dokon_telefon", "+998 XX XXX XX XX")
    dokon_inn = data.get("dokon_inn", "")
    dokon_manzil = data.get("dokon_manzil", "")
    klient   = data.get("klient_ismi", "")
    klient_tel = data.get("klient_telefon", "")
    klient_manzil = data.get("klient_manzil", "")
    klient_inn = data.get("klient_inn", "")
    tovarlar = data.get("tovarlar", [])
    jami     = data.get("jami_summa", 0)
    qarz     = data.get("qarz", 0)
    tolangan = data.get("tolangan", jami)

    def _shading(cell, color):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), color)
        tc_pr.append(shd)

    # ═══ 1. SARLAVHA — Do'kon nomi ═══
    h = doc.add_heading("", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run(dokon)
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = KOK

    # Do'kon manzili
    if dokon_manzil:
        p_addr = doc.add_paragraph()
        p_addr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_addr.add_run(dokon_manzil).font.size = Pt(10)

    # Do'kon tel + INN
    tel_inn_parts = []
    if dokon_tel: tel_inn_parts.append(f"Tel: {dokon_tel}")
    if dokon_inn: tel_inn_parts.append(f"INN: {dokon_inn}")
    if tel_inn_parts:
        p_ti = doc.add_paragraph()
        p_ti.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_ti.add_run("  ".join(tel_inn_parts)).font.size = Pt(10)

    doc.add_paragraph()

    # ═══ 2. NAKLADNOY sarlavha ═══
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_title.add_run("NAKLADNOY")
    run_t.font.size = Pt(16)
    run_t.font.bold = True
    run_t.font.color.rgb = KOK

    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_info.add_run(f"#{inv_no}   Sana: {sana}").font.size = Pt(11)

    doc.add_paragraph()

    # ═══ 3. MIJOZ MA'LUMOTLARI jadvali ═══
    mijoz_tbl = doc.add_table(rows=5, cols=2)
    mijoz_tbl.style = "Table Grid"

    # Sarlavha qator (birlashtirilgan)
    mijoz_tbl.cell(0, 0).merge(mijoz_tbl.cell(0, 1))
    cell_h = mijoz_tbl.cell(0, 0)
    cell_h.text = "MIJOZ MA'LUMOTLARI"
    for p in cell_h.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    _shading(cell_h, "1a56db")

    # Ma'lumot qatorlari
    rows_data = [
        ("Kompaniya:", klient or "—"),
        ("Manzil:", klient_manzil or "Mijoz manzili"),
        ("Telefon:", klient_tel or "+998 XX XXX XX XX"),
        ("INN:", klient_inn or "XXXXXXXXX"),
    ]
    for i, (label, value) in enumerate(rows_data, 1):
        mijoz_tbl.cell(i, 0).text = label
        mijoz_tbl.cell(i, 1).text = value
        for cell in [mijoz_tbl.cell(i, 0), mijoz_tbl.cell(i, 1)]:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
        # Label qalin
        for p in mijoz_tbl.cell(i, 0).paragraphs:
            for r in p.runs:
                r.font.bold = True

    doc.add_paragraph()

    # ═══ 4. TOVARLAR JADVALI ═══
    sarlavhalar = ["№", "Tovar/Xizmat nomi", "Birlik", "Miqdor", "Narx", "Summa"]
    tv = doc.add_table(rows=1 + len(tovarlar) + 1, cols=6)
    tv.style = "Table Grid"

    # Sarlavha qator
    for ci, s in enumerate(sarlavhalar):
        cell = tv.cell(0, ci)
        cell.text = s
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shading(cell, "1a56db")

    # Tovar qatorlari
    jami_hisob = Decimal("0")
    for i, t in enumerate(tovarlar, 1):
        miq = Decimal(str(t.get("miqdor", 0)))
        bir = t.get("birlik", "dona")
        narx = Decimal(str(t.get("narx", 0)))
        tj = Decimal(str(t.get("jami", 0))) or miq * narx
        jami_hisob += tj

        miq_s = f"{miq:.1f}".rstrip("0").rstrip(".")
        qiymatlar = [
            str(i),
            t.get("nomi", ""),
            bir,
            miq_s,
            _pul(narx) if narx else "—",
            _pul(tj),
        ]
        soya = i % 2 == 0
        for ci, qiy in enumerate(qiymatlar):
            cell = tv.cell(i, ci)
            cell.text = qiy
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
            if soya:
                _shading(cell, "f7f9fc")

    # JAMI qator
    oxirgi = len(tovarlar) + 1
    tv.cell(oxirgi, 0).merge(tv.cell(oxirgi, 4))
    tv.cell(oxirgi, 0).text = ""
    tv.cell(oxirgi, 4).text = ""

    # JAMI label
    jami_cell_label = tv.cell(oxirgi, 0)
    for p in jami_cell_label.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_j = jami_cell_label.paragraphs[0].add_run("JAMI:")
    run_j.font.bold = True
    run_j.font.size = Pt(11)

    # JAMI summa
    jami_val = Decimal(str(jami)) if jami else jami_hisob
    jami_cell = tv.cell(oxirgi, 5)
    jami_cell.text = ""
    run_js = jami_cell.paragraphs[0].add_run(_pul(jami_val))
    run_js.font.bold = True
    run_js.font.size = Pt(11)

    doc.add_paragraph()

    # Qarz
    if float(qarz or 0) > 0:
        q_par = doc.add_paragraph()
        q_par.add_run(f"✅ To'landi: {_pul(tolangan)} so'm     ")
        r = q_par.add_run(f"⚠️ QARZ: {_pul(qarz)} so'm")
        r.font.bold = True
        r.font.color.rgb = QIZIL

    doc.add_paragraph()

    # ═══ 5. IMZO VA MUHR ═══
    imzo_tbl = doc.add_table(rows=4, cols=2)
    imzo_tbl.style = "Table Grid"

    imzo_data = [
        (f"TOPSHIRDI: ________________", f"QABUL QILDI: ________________"),
        (f"F.I.O: ________________", f"F.I.O: ________________"),
        (f"Imzo: ________________", f"Imzo: ________________"),
        (f"Sana: {sana}", f"Sana: {sana}"),
    ]
    for ri, (left, right) in enumerate(imzo_data):
        imzo_tbl.cell(ri, 0).text = left
        imzo_tbl.cell(ri, 1).text = right
        for ci in range(2):
            for p in imzo_tbl.cell(ri, ci).paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    if ri == 0:
                        r.font.bold = True

    doc.add_paragraph()

    # Footer
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot.add_run(f"@mashrab_ceo_bot  |  {dokon}  |  {_hozir_full()}").font.size = Pt(8)
    foot.runs[0].font.color.rgb = KULRANG

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ════════════ EXCEL ════════════

def yaratish_excel(data: dict) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook(); ws = wb.active; ws.title = "Nakladnoy"
    KOK = "1a56db"; OCH = "f7f9fc"; QIZIL = "dc2626"; OQ = "FFFFFF"

    def _sh(qalin=False, rang="000000", o=10):
        return Font(bold=qalin, color=rang, size=o, name="Calibri")
    def _tf(rang): return PatternFill("solid", fgColor=rang)
    def _ch():
        s = Side(style="thin", color="d1d5db")
        return Border(left=s, right=s, top=s, bottom=s)

    inv_no   = data.get("invoice_number", "N/A")
    sana     = data.get("date", _hozir())
    dokon    = data.get("dokon_nomi", "Mashrab Moliya")
    klient   = data.get("klient_ismi", "")
    klient_tel = data.get("klient_telefon", "")
    klient_manzil = data.get("klient_manzil", "")
    klient_inn = data.get("klient_inn", "")
    tovarlar = data.get("tovarlar", [])
    jami     = data.get("jami_summa", 0)
    qarz     = data.get("qarz", 0)
    tolangan = data.get("tolangan", jami)

    # Sarlavha
    ws["A1"] = dokon
    ws["A1"].font = _sh(qalin=True, rang=KOK, o=16)
    ws.merge_cells("A1:F1")
    ws["A2"] = f"NAKLADNOY #{inv_no}   |   Sana: {sana}"
    ws["A2"].alignment = Alignment(horizontal="center")
    ws.merge_cells("A2:F2")

    # Mijoz ma'lumotlari
    ws["A4"] = "MIJOZ MA'LUMOTLARI"
    ws["A4"].font = _sh(qalin=True, rang=OQ)
    ws["A4"].fill = _tf(KOK)
    ws.merge_cells("A4:F4")

    mijoz_rows = [
        ("Kompaniya:", klient or "—"),
        ("Manzil:", klient_manzil or "—"),
        ("Telefon:", klient_tel or "—"),
        ("INN:", klient_inn or "—"),
    ]
    for i, (label, val) in enumerate(mijoz_rows, 5):
        ws.cell(i, 1, label).font = _sh(qalin=True)
        ws.cell(i, 2, val)

    # Tovar jadvali
    start_row = 10
    sarlavhalar = ["№", "Tovar/Xizmat nomi", "Birlik", "Miqdor", "Narx", "Summa"]
    kengliklar = [5, 30, 10, 10, 15, 18]
    for ci, (s, k) in enumerate(zip(sarlavhalar, kengliklar), 1):
        cell = ws.cell(row=start_row, column=ci, value=s)
        cell.font = _sh(qalin=True, rang=OQ)
        cell.fill = _tf(KOK)
        cell.alignment = Alignment(horizontal="center")
        cell.border = _ch()
        ws.column_dimensions[get_column_letter(ci)].width = k

    jami_hisob = Decimal("0")
    for i, t in enumerate(tovarlar, 1):
        miq = Decimal(str(t.get("miqdor", 0)))
        narx = Decimal(str(t.get("narx", 0)))
        tj = Decimal(str(t.get("jami", 0))) or miq * narx
        jami_hisob += tj
        miq_s = f"{miq:.1f}".rstrip("0").rstrip(".")
        qiymatlar = [i, t.get("nomi", ""), t.get("birlik", "dona"),
                     miq_s, str(narx) if narx else "—", str(tj)]
        soya = i % 2 == 0
        for ci, qiy in enumerate(qiymatlar, 1):
            cell = ws.cell(row=start_row + i, column=ci, value=qiy)
            cell.font = _sh()
            cell.border = _ch()
            if soya: cell.fill = _tf(OCH)
            if ci in (5, 6): cell.alignment = Alignment(horizontal="right")

    # JAMI
    total_row = start_row + len(tovarlar) + 1
    ws.cell(total_row, 5, "JAMI:").font = _sh(qalin=True)
    jami_val = Decimal(str(jami)) if jami else jami_hisob
    ws.cell(total_row, 6, str(jami_val)).font = _sh(qalin=True)

    if float(qarz or 0) > 0:
        ws.cell(total_row + 1, 5, "To'landi:").font = _sh(qalin=True)
        ws.cell(total_row + 1, 6, str(tolangan))
        ws.cell(total_row + 2, 5, "QARZ:").font = _sh(qalin=True, rang=QIZIL)
        ws.cell(total_row + 2, 6, str(qarz)).font = _sh(qalin=True, rang=QIZIL)

    # Imzo
    sig_row = total_row + 4
    ws.cell(sig_row, 1, "TOPSHIRDI: ________________").font = _sh(qalin=True)
    ws.cell(sig_row, 4, "QABUL QILDI: ________________").font = _sh(qalin=True)
    ws.cell(sig_row + 1, 1, "F.I.O: ________________")
    ws.cell(sig_row + 1, 4, "F.I.O: ________________")
    ws.cell(sig_row + 2, 1, "Imzo: ________________")
    ws.cell(sig_row + 2, 4, "Imzo: ________________")
    ws.cell(sig_row + 3, 1, f"Sana: {sana}")
    ws.cell(sig_row + 3, 4, f"Sana: {sana}")

    ws.cell(sig_row + 5, 1, f"@mashrab_ceo_bot | {_hozir_full()}").font = Font(
        size=8, color="6b7280", name="Calibri"
    )

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ════════════ UCHALA FORMAT ════════════

def uchala_format(data: dict) -> dict[str, bytes]:
    """
    Word, Excel va PDF nakladnoy yaratish.
    Qaytaradi: {"word": bytes, "excel": bytes, "pdf": bytes}
    """
    from services.bot.bot_services.export_pdf import sotuv_pdf

    natija: dict[str, bytes] = {}
    try:
        natija["word"] = yaratish_word(data)
        log.info("✅ Nakladnoy Word OK")
    except Exception as e:
        log.error("Nakladnoy Word xato: %s", e, exc_info=True)
    try:
        natija["excel"] = yaratish_excel(data)
        log.info("✅ Nakladnoy Excel OK")
    except Exception as e:
        log.error("Nakladnoy Excel xato: %s", e, exc_info=True)
    try:
        pdf_data = {
            "klient":     data.get("klient_ismi") or data.get("klient"),
            "tovarlar":   data.get("tovarlar", []),
            "jami_summa": data.get("jami_summa", 0),
            "qarz":       data.get("qarz", 0),
            "tolandan":   data.get("tolangan", 0),
            "manba":      None,
        }
        natija["pdf"] = sotuv_pdf(pdf_data, data.get("dokon_nomi", "Mashrab Moliya"))
        log.info("✅ Nakladnoy PDF OK")
    except Exception as e:
        log.error("Nakladnoy PDF xato: %s", e, exc_info=True)
    return natija
