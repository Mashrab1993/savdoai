"""
╔══════════════════════════════════════════════════════════════╗
║  MASHRAB MOLIYA v21.3 — CELERY WORKER                         ║
║  Og'ir ishlarni orqada bajarish:                            ║
║  ✅ Kunlik/haftalik hisobot                                  ║
║  ✅ Qarz eslatmasi (SMS/Telegram)                           ║
║  ✅ Nakladnoy generatsiya                                    ║
║  ✅ Katta export (Excel 1000+ qator)                        ║
║  ✅ DB backup (kunlik)                                      ║
╚══════════════════════════════════════════════════════════════╝
"""
from __future__ import annotations
import logging
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.dirname(__file__))))

from celery import Celery
from celery.schedules import crontab

log = logging.getLogger(__name__)

REDIS_URL = os.getenv("REDIS_URL", "redis://localhost:6379/0")

app = Celery(
    "mashrab_moliya",
    broker=REDIS_URL,
    backend=REDIS_URL,
)

app.conf.update(
    task_serializer          = "json",
    accept_content           = ["json"],
    result_serializer        = "json",
    timezone                 = "Asia/Tashkent",
    enable_utc               = True,
    task_track_started       = True,
    task_soft_time_limit     = 300,
    task_time_limit          = 600,
    worker_prefetch_multiplier = 1,
    task_acks_late           = True,
    # Exponential backoff retry
    task_max_retries         = 3,
    task_default_retry_delay = 60,
    # Dead Letter Queue (muvaffaqiyatsiz tasklar)
    task_reject_on_worker_lost = True,
    beat_schedule={
        # Kunlik hisobot — 22:00
        "kunlik-hisobot": {
            "task":     "tasks.kunlik_hisobot_barcha",
            "schedule": crontab(hour=22, minute=0),
        },
        # Haftalik hisobot — Dushanba 08:00
        "haftalik-hisobot": {
            "task":     "tasks.haftalik_hisobot_barcha",
            "schedule": crontab(hour=8, minute=0, day_of_week=1),
        },
        # Qarz eslatmasi — 10:00
        "qarz-eslatma": {
            "task":     "tasks.qarz_eslatma_barcha",
            "schedule": crontab(hour=10, minute=0),
        },
        # DB backup — Juma 03:00
        "db-backup": {
            "task":     "tasks.db_backup",
            "schedule": crontab(hour=3, minute=0, day_of_week=5),
        },
        # Obuna eslatmasi — 09:00
        "obuna-eslatma": {
            "task":     "tasks.obuna_eslatma_barcha",
            "schedule": crontab(hour=9, minute=0),
        },
        # Ledger reconciliation — har kuni 06:00
        "ledger-recon": {
            "task":     "tasks.ledger_reconciliation",
            "schedule": crontab(hour=6, minute=0),
        },
    },
)


@app.task(bind=True, name="tasks.kunlik_hisobot_barcha",
          max_retries=3, default_retry_delay=60)
def kunlik_hisobot_barcha(self):
    """Barcha faol foydalanuvchilarga kunlik hisobot yuborish"""
    import asyncio
    try:
        asyncio.run(_kunlik_hisobot_async())
    except Exception as exc:
        log.error("Kunlik hisobot xato: %s", exc)
        raise self.retry(exc=exc)


async def _kunlik_hisobot_async():
    """Barcha faol userlarga bugungi sotuv hisobotini yuborish."""
    from shared.database.pool import pool_init, get_pool, rls_conn
    dsn       = os.environ["DATABASE_URL"]
    bot_token = os.environ.get("BOT_TOKEN", "")
    await pool_init(dsn, min_size=1, max_size=3)

    async with get_pool().acquire() as c:
        users = await c.fetch("SELECT id FROM users WHERE faol=TRUE")

    if not bot_token:
        log.warning("BOT_TOKEN yo'q — kunlik hisobot yuborilmadi")
        return

    import httpx
    async with httpx.AsyncClient(timeout=10) as http:
        for user in users:
            try:
                async with rls_conn(user["id"]) as c:
                    bugun = await c.fetchrow("""
                        SELECT
                            COUNT(ss.id)              AS soni,
                            COALESCE(SUM(ss.jami), 0) AS jami,
                            COALESCE(SUM(ss.qarz), 0) AS qarz
                        FROM sotuv_sessiyalar ss
                        WHERE (ss.sana AT TIME ZONE 'Asia/Tashkent')::date = CURRENT_DATE
                    """)
                    jami_qarz = await c.fetchval(
                        "SELECT COALESCE(SUM(qolgan),0) FROM qarzlar WHERE yopildi=FALSE"
                    )

                if not bugun or bugun["soni"] == 0:
                    continue  # bo'sh kun — xabar yuborma

                matn = (
                    "📊 *Bugungi hisobot*\n"
                    f"Sotuv: {bugun['soni']} ta — {bugun['jami']:,} so'm\n"
                    f"Yangi qarz: {bugun['qarz']:,} so'm\n"
                    f"Jami qarz: {float(jami_qarz or 0):,.0f} so'm"
                )
                await http.post(
                    f"https://api.telegram.org/bot{bot_token}/sendMessage",
                    json={"chat_id": user["id"], "text": matn, "parse_mode": "Markdown"},
                )
                log.info("Kunlik hisobot yuborildi uid=%d", user["id"])
            except Exception as e:
                log.warning("Kunlik hisobot uid=%d xato: %s", user["id"], e)


@app.task(bind=True, name="tasks.haftalik_hisobot_barcha",
          max_retries=3, default_retry_delay=60)
def haftalik_hisobot_barcha(self):
    import asyncio
    try:
        asyncio.run(_haftalik_hisobot_async())
    except Exception as exc:
        raise self.retry(exc=exc)


async def _haftalik_hisobot_async():
    from shared.database.pool import pool_init, get_pool, rls_conn
    dsn = os.environ["DATABASE_URL"]
    bot_token = os.environ.get("BOT_TOKEN", "")
    await pool_init(dsn, min_size=1, max_size=3)

    async with get_pool().acquire() as c:
        users = await c.fetch("SELECT id FROM users WHERE faol=TRUE")

    if not bot_token:
        log.warning("BOT_TOKEN yo'q — haftalik hisobot yuborilmadi")
        return

    import httpx
    async with httpx.AsyncClient(timeout=10) as http:
        for user in users:
            try:
                async with rls_conn(user["id"]) as c:
                    h = await c.fetchrow("""
                        SELECT
                            COUNT(ss.id)              AS soni,
                            COALESCE(SUM(ss.jami), 0) AS jami,
                            COALESCE(SUM(ss.qarz), 0) AS qarz
                        FROM sotuv_sessiyalar ss
                        WHERE ss.sana >= NOW() - INTERVAL '7 days'
                    """)
                    jami_qarz = await c.fetchval(
                        "SELECT COALESCE(SUM(qolgan),0) FROM qarzlar WHERE yopildi=FALSE"
                    )

                if not h or h["soni"] == 0:
                    continue  # bo'sh hafta — xabar yuborma

                matn = (
                    f"📊 *Haftalik hisobot (7 kun)*\n"
                    f"Sotuv: {h['soni']} ta — {float(h['jami']):,.0f} so'm\n"
                    f"Yangi qarz: {float(h['qarz']):,.0f} so'm\n"
                    f"Jami qarz: {float(jami_qarz or 0):,.0f} so'm"
                )
                await http.post(
                    f"https://api.telegram.org/bot{bot_token}/sendMessage",
                    json={"chat_id": user["id"], "text": matn, "parse_mode": "Markdown"},
                )
                log.info("Haftalik hisobot yuborildi uid=%d", user["id"])
            except Exception as e:
                log.warning("Haftalik uid=%d xato: %s", user["id"], e)


@app.task(bind=True, name="tasks.qarz_eslatma_barcha",
          max_retries=2)
def qarz_eslatma_barcha(self):
    import asyncio
    try:
        asyncio.run(_qarz_eslatma_async())
    except Exception as exc:
        raise self.retry(exc=exc)


async def _qarz_eslatma_async():
    from shared.database.pool import pool_init, get_pool, rls_conn
    dsn = os.environ["DATABASE_URL"]
    bot_token = os.environ.get("BOT_TOKEN", "")
    await pool_init(dsn, min_size=1, max_size=3)

    async with get_pool().acquire() as c:
        users = await c.fetch("SELECT id FROM users WHERE faol=TRUE")

    if not bot_token:
        log.warning("BOT_TOKEN yo'q — qarz eslatma yuborilmadi")
        return

    import httpx
    async with httpx.AsyncClient(timeout=10) as http:
        for user in users:
            try:
                async with rls_conn(user["id"]) as c:
                    qarzlar = await c.fetch("""
                        SELECT klient_ismi, SUM(qolgan) AS jami
                        FROM qarzlar
                        WHERE yopildi=FALSE AND qolgan > 0
                          AND (muddat IS NULL OR muddat <= CURRENT_DATE + 3)
                        GROUP BY klient_ismi
                        ORDER BY jami DESC
                        LIMIT 5
                    """)

                if not qarzlar:
                    continue

                lines = [f"⚠️ *Bugun muddat o'tgan/yaqin qarzlar*"]
                for q in qarzlar:
                    lines.append(f"• {q['klient_ismi']}: {float(q['jami']):,.0f} so'm")
                matn = "\n".join(lines)

                await http.post(
                    f"https://api.telegram.org/bot{bot_token}/sendMessage",
                    json={"chat_id": user["id"], "text": matn, "parse_mode": "Markdown"},
                )
                log.info("Qarz eslatma yuborildi uid=%d: %d klient", user["id"], len(qarzlar))
            except Exception as e:
                log.warning("Qarz eslatma uid=%d xato: %s", user["id"], e)


@app.task(bind=True, name="tasks.obuna_eslatma_barcha",
          max_retries=2, default_retry_delay=120)
def obuna_eslatma_barcha(self):
    import asyncio
    try:
        asyncio.run(_obuna_eslatma_async())
    except Exception as exc:
        log.error("Obuna eslatma task xato: %s", exc)
        raise self.retry(exc=exc)


async def _obuna_eslatma_async():
    """Obuna 3 kunda tugayotgan foydalanuvchilarga eslatma yuborish."""
    from shared.database.pool import pool_init, get_pool
    dsn       = os.environ["DATABASE_URL"]
    bot_token = os.environ.get("BOT_TOKEN", "")
    await pool_init(dsn, min_size=1, max_size=2)

    async with get_pool().acquire() as c:
        tugayotganlar = await c.fetch("""
            SELECT id, ism, obuna_tugash
            FROM users
            WHERE faol=TRUE
              AND obuna_tugash IS NOT NULL
              AND obuna_tugash = CURRENT_DATE + 3
        """)

    log.info("Obuna eslatma: %d foydalanuvchi", len(tugayotganlar))
    if not tugayotganlar:
        return
    if not bot_token:
        log.warning("BOT_TOKEN yo'q — obuna eslatma yuborilmadi")
        return

    import httpx
    async with httpx.AsyncClient(timeout=10) as http:
        for u in tugayotganlar:
            try:
                matn = (
                    f"⚠️ Hurmatli {u['ism'] or 'foydalanuvchi'},\n"
                    f"Sizning obunangiz *{u['obuna_tugash']}* da tugaydi.\n"
                    f"Uzluksiz ishlash uchun obunani yangilang."
                )
                await http.post(
                    f"https://api.telegram.org/bot{bot_token}/sendMessage",
                    json={"chat_id": u["id"], "text": matn, "parse_mode": "Markdown"},
                )
                log.info("Obuna eslatma yuborildi uid=%d", u["id"])
            except Exception as e:
                log.warning("Obuna eslatma uid=%d xato: %s", u["id"], e)


@app.task(name="tasks.katta_export")
def katta_export(user_id: int, export_turi: str,
                  sana_dan: str, sana_gacha: str,
                  format_: str = "excel") -> dict:
    """Katta Excel/PDF export (background). format_: excel | pdf"""
    import asyncio
    try:
        natija = asyncio.run(
            _export_async(user_id, export_turi, sana_dan, sana_gacha, format_)
        )
        # Faylni base64 ga o'girib Redis da saqlash (cross-container delivery)
        import base64, os as _os
        MAX_BYTES = 8 * 1024 * 1024  # 8 MB limit
        file_size = _os.path.getsize(natija) if _os.path.exists(natija) else 0
        if file_size > MAX_BYTES:
            log.warning("Export fayl hajmi katta: %d bytes (uid=%d)", file_size, user_id)
            return {
                "status":  "katta_fayl",
                "format":  format_,
                "hajm_kb": file_size // 1024,
                "xato":    f"Fayl hajmi {file_size//1024}KB — 8MB dan oshmasligi kerak",
            }
        with open(natija, "rb") as f:
            content_b64 = base64.b64encode(f.read()).decode("ascii")
        # Vaqtinchalik faylni o'chirish
        try: _os.unlink(natija)
        except Exception as _e: log.debug("silent: %s", _e)
        return {
            "status":      "tayyor",
            "format":      format_,
            "user_id":     user_id,
            "content_b64": content_b64,
            "hajm_kb":     file_size // 1024,
        }
    except Exception as e:
        log.error("Export xato uid=%d format=%s: %s", user_id, format_, e)
        return {"status": "xato", "xato": "Export bajarilmadi. Keyinroq urinib ko'ring."}


async def _export_async(user_id: int, export_turi: str,
                         sana_dan: str, sana_gacha: str,
                         format_: str = "excel") -> str:
    """
    Real export yaratish.
    format_: excel (xlsx) yoki pdf
    Fayl: /tmp/export_<uid>_<dan>_<gacha>.<ext>
    """
    from shared.database.pool import pool_init, rls_conn
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment
    except ImportError:
        raise RuntimeError("openpyxl o'rnatilmagan (pip install openpyxl)")

    dsn = os.environ["DATABASE_URL"]
    await pool_init(dsn, min_size=1, max_size=2)

    # Sanani export_turi dan hisoblash (sana_dan/sana_gacha ustunlik qiladi)
    import datetime as _dt
    today = _dt.date.today()
    if not sana_dan:
        if export_turi == "haftalik":
            sana_dan = str(today - _dt.timedelta(days=7))
        elif export_turi == "oylik":
            sana_dan = str(today - _dt.timedelta(days=30))
        else:  # kunlik (default)
            sana_dan = str(today)
    if not sana_gacha:
        sana_gacha = str(today)

    async with rls_conn(user_id) as c:
        # Sotuv sessiyalari
        rows = await c.fetch("""
            SELECT ss.id, ss.klient_ismi, ss.jami, ss.tolangan, ss.qarz,
                   ss.sana, ss.izoh
            FROM sotuv_sessiyalar ss
            WHERE ss.sana::DATE BETWEEN $1::DATE AND $2::DATE
            ORDER BY ss.sana DESC
        """, sana_dan, sana_gacha)

        # Foydalanuvchi nomi
        user_row = await c.fetchrow("SELECT ism, dokon_nomi FROM users WHERE id=$1", user_id)
        dokon = (user_row["dokon_nomi"] if user_row else "Do'kon") or "Do'kon"

    # Excel yaratish
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sotuv"

    # Sarlavha
    ws["A1"] = f"{dokon} — Sotuv hisoboti"
    ws["A1"].font = Font(bold=True, size=14)
    ws["A2"] = f"Davr: {sana_dan or 'boshi'} → {sana_gacha or 'bugun'}"
    ws["A3"] = f"Yaratildi: {__import__('datetime').datetime.now().strftime('%Y-%m-%d %H:%M')}"

    # Ustun sarlavhalari
    headers = ["#", "Sana", "Klient", "Jami (so'm)", "To'langan", "Qarz", "Izoh"]
    header_row = 5
    header_fill = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=header_row, column=col, value=h)
        cell.font  = Font(bold=True, color="FFFFFF")
        cell.fill  = header_fill
        cell.alignment = Alignment(horizontal="center")

    # Ma'lumotlar
    jami_total = 0
    for r_idx, row in enumerate(rows, 1):
        data_row = header_row + r_idx
        ws.cell(data_row, 1, r_idx)
        ws.cell(data_row, 2, row["sana"].strftime("%Y-%m-%d") if hasattr(row["sana"], "strftime") else str(row["sana"]))
        ws.cell(data_row, 3, row["klient_ismi"] or "—")
        ws.cell(data_row, 4, float(row["jami"] or 0))
        ws.cell(data_row, 5, float(row["tolangan"] or 0))
        qarz = float(row["qarz"] or 0)
        cell_qarz = ws.cell(data_row, 6, qarz)
        if qarz > 0:
            cell_qarz.fill = PatternFill(start_color="FFCCCC", end_color="FFCCCC", fill_type="solid")
        ws.cell(data_row, 7, row["izoh"] or "")
        jami_total += float(row["jami"] or 0)

    # Jami qator
    total_row = header_row + len(rows) + 1
    ws.cell(total_row, 3, "JAMI:").font = Font(bold=True)
    cell_total = ws.cell(total_row, 4, jami_total)
    cell_total.font = Font(bold=True)

    # Ustun kengliklari
    col_widths = [5, 12, 25, 16, 16, 14, 30]
    for i, width in enumerate(col_widths, 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = width

    ext  = "xlsx" if format_ == "excel" else "pdf"
    fayl = f"/tmp/export_{user_id}_{sana_dan or 'all'}_{sana_gacha or 'all'}.{ext}"

    if format_ == "excel":
        wb.save(fayl)
    elif format_ == "pdf":
        # PDF: export same data via reportlab
        fayl = await _export_pdf_async(user_id, rows, dokon, sana_dan, sana_gacha)
    else:
        wb.save(fayl)  # fallback to excel

    log.info("✅ Export: %d qator, format=%s, fayl=%s (%.1fKB)",
             len(rows), format_, fayl, __import__("os").path.getsize(fayl)/1024)
    return fayl



async def _export_pdf_async(user_id: int, rows, dokon: str,
                              sana_dan: str, sana_gacha: str) -> str:
    """Sotuv hisobotini PDF formatda yaratish (reportlab)."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except ImportError:
        raise RuntimeError("reportlab o'rnatilmagan (pip install reportlab)")

    import datetime
    fayl = f"/tmp/export_{user_id}_{sana_dan or 'all'}_{sana_gacha or 'all'}.pdf"
    doc  = SimpleDocTemplate(fayl, pagesize=A4)
    styles = getSampleStyleSheet()
    elements = []

    # Sarlavha
    elements.append(Paragraph(f"{dokon} — Sotuv hisoboti", styles['Title']))
    elements.append(Paragraph(f"Davr: {sana_dan or 'boshi'} — {sana_gacha or 'bugun'}", styles['Normal']))
    elements.append(Paragraph(f"Yaratildi: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}", styles['Normal']))
    elements.append(Spacer(1, 12))

    # Jadval sarlavhasi
    table_data = [["#", "Sana", "Klient", "Jami (so'm)", "Qarz"]]
    for i, row in enumerate(rows, 1):
        sana_str = row["sana"].strftime("%Y-%m-%d") if hasattr(row["sana"], "strftime") else str(row["sana"])
        table_data.append([
            str(i),
            sana_str,
            row["klient_ismi"] or "—",
            f"{float(row['jami'] or 0):,.0f}",
            f"{float(row['qarz'] or 0):,.0f}",
        ])

    # Jami
    jami_total = sum(float(r["jami"] or 0) for r in rows)
    table_data.append(["", "", "JAMI:", f"{jami_total:,.0f}", ""])

    tbl = Table(table_data, colWidths=[30, 80, 160, 100, 80])
    tbl.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1F4E79')),
        ('TEXTCOLOR',  (0, 0), (-1, 0), colors.white),
        ('FONTNAME',   (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('GRID',       (0, 0), (-1, -1), 0.5, colors.grey),
        ('ROWBACKGROUNDS', (0, 1), (-1, -2), [colors.white, colors.HexColor('#F2F2F2')]),
        ('FONTNAME',   (0, -1), (-1, -1), 'Helvetica-Bold'),
    ]))
    elements.append(tbl)
    doc.build(elements)
    return fayl

async def _yuborish(http, bot_token: str,
                    user_id: int, tur: str):
    """Telegram xabar yuborish"""
    url = f"https://api.telegram.org/bot{bot_token}/sendMessage"
    matn = f"📊 {tur.capitalize()} hisobot tayyor!"
    await http.post(url, json={
        "chat_id": user_id,
        "text": matn,
        "parse_mode": "Markdown",
    }, timeout=10)


@app.task(bind=True, name="tasks.db_backup",
          max_retries=2, default_retry_delay=300)
def db_backup(self):
    """
    Avtomatik DB backup (haftalik).
    pg_dump → /tmp/backup_YYYYMMDD.sql.gz
    Ishlab chiqarishda: S3/R2 ga yuklash kerak.
    """
    import subprocess, gzip, datetime as dt
    try:
        sana    = dt.datetime.now().strftime("%Y%m%d_%H%M")
        out_gz  = f"/tmp/mm_backup_{sana}.sql.gz"
        db_url  = os.environ["DATABASE_URL"]

        result = subprocess.run(
            ["pg_dump", "--no-owner", "--no-privileges", db_url],
            capture_output=True, timeout=300
        )
        if result.returncode != 0:
            raise RuntimeError(f"pg_dump xato: {result.stderr.decode()}")

        with gzip.open(out_gz, "wb") as f:
            f.write(result.stdout)

        size_kb = os.path.getsize(out_gz) // 1024
        log.info("✅ DB backup: %s (%d KB)", out_gz, size_kb)

        # TODO: S3/R2 ga yuklash
        # boto3.upload_file(out_gz, "mashrab-backups", f"backup_{sana}.sql.gz")

        return {"status": "ok", "fayl": out_gz, "hajm_kb": size_kb}
    except Exception as exc:
        log.error("DB backup xato: %s", exc)
        raise self.retry(exc=exc)


@app.task(name="tasks.nakladnoy_yaratish")
def nakladnoy_yaratish(user_id: int, sessiya_id: int,
                        format_: str = "word") -> dict:
    """
    Nakladnoy hujjat yaratish va Telegram orqali yuborish.
    format_: word | excel | pdf
    Bot orqali chaqirilganda: result faylni bot foydalanuvchiga yuboradi.
    """
    import asyncio
    try:
        fayl = asyncio.run(_nakl_async(user_id, sessiya_id, format_))
        # Telegram orqali yetkazish
        bot_token = os.environ.get("BOT_TOKEN", "")
        if bot_token and os.path.exists(fayl):
            asyncio.run(_nakl_yuborish(bot_token, user_id, fayl, sessiya_id))
        # content_b64 saqlash (API delivery uchun)
        import base64, os as _os2
        content_b64 = ""
        if _os2.path.exists(fayl):
            with open(fayl, "rb") as f:
                content_b64 = base64.b64encode(f.read()).decode("ascii")
            try: _os2.unlink(fayl)
            except Exception as _e: log.debug("silent: %s", _e)
        if not content_b64:
            log.error("Nakladnoy: content_b64 bo'sh — fayl yaratilmadi uid=%d sess=%d",
                      user_id, sessiya_id)
            return {
                "status":    "xato",
                "xato":      "Nakladnoy yaratilmadi (fayl bo'sh)",
                "user_id":   user_id,
                "sessiya_id": sessiya_id,
            }
        return {
            "status":      "tayyor",
            "format":      format_,
            "user_id":     user_id,
            "sessiya_id":  sessiya_id,
            "content_b64": content_b64,
        }
    except Exception as e:
        log.error("Nakladnoy uid=%d sess=%d: %s", user_id, sessiya_id, e, exc_info=True)
        return {"status": "xato", "xato": "Nakladnoy yaratilmadi"}



async def _nakl_yuborish(bot_token: str, user_id: int,
                          fayl: str, sessiya_id: int) -> None:
    """Nakladnoy faylni foydalanuvchiga Telegram orqali yuborish."""
    import httpx
    ext = fayl.rsplit(".", 1)[-1] if "." in fayl else "docx"
    fname = f"Nakladnoy_{sessiya_id}.{ext}"
    if not os.path.exists(fayl):
        log.warning("Nakladnoy delivery: fayl topilmadi %s", fayl)
        return
    try:
        async with httpx.AsyncClient(timeout=30) as http:
            with open(fayl, "rb") as f:
                await http.post(
                    f"https://api.telegram.org/bot{bot_token}/sendDocument",
                    data={"chat_id": user_id, "caption": f"📋 Nakladnoy #{sessiya_id}"},
                    files={"document": (fname, f)},
                )
        log.info("✅ Nakladnoy yuborildi uid=%d sess=%d", user_id, sessiya_id)
    except Exception as e:
        log.warning("Nakladnoy yubormadi uid=%d: %s", user_id, e)

async def _nakl_async(user_id: int, sessiya_id: int,
                       format_: str) -> str:
    """
    Nakladnoy hujjat yaratish.
    format_: word (docx) | excel (xlsx) | pdf
    """
    from shared.database.pool import pool_init, rls_conn

    dsn = os.environ["DATABASE_URL"]
    await pool_init(dsn, min_size=1, max_size=2)

    async with rls_conn(user_id) as c:
        sess = await c.fetchrow(
            "SELECT id, klient_ismi, jami, tolangan, qarz, sana FROM sotuv_sessiyalar WHERE id=$1 AND user_id=$2", sessiya_id, user_id
        )
        if not sess:
            raise ValueError(f"Sessiya {sessiya_id} topilmadi")
        tovarlar = await c.fetch(
            "SELECT tovar_nomi, miqdor, birlik, sotish_narxi, jami FROM chiqimlar WHERE sessiya_id=$1 AND user_id=$2", sessiya_id, user_id
        )
        nakl_no = await c.fetchval(
            "SELECT nakladnoy_raqami_ol($1)", user_id
        )

    data = {
        "nakl_no":    nakl_no,
        "sessiya_id": sessiya_id,
        "klient":     sess["klient_ismi"] or "",
        "jami_summa": float(sess["jami"]),
        "tolangan":   float(sess["tolangan"]),
        "qarz":       float(sess["qarz"]),
        "tovarlar": [
            {
                "nomi":     t["tovar_nomi"],
                "miqdor":   float(t["miqdor"]),
                "birlik":   t["birlik"],
                "narx":     float(t["sotish_narxi"]),
                "jami":     float(t["jami"]),
            }
            for t in tovarlar
        ],
    }

    ext = {"word": "docx", "excel": "xlsx", "pdf": "pdf"}.get(format_, "docx")
    fayl = f"/tmp/nakl_{user_id}_{sessiya_id}.{ext}"

    if format_ == "excel":
        _nakl_excel(data, fayl)
    elif format_ == "pdf":
        _nakl_pdf(data, fayl)
    else:
        _nakl_word(data, fayl)

    log.info("✅ Nakladnoy: uid=%d sess=%d format=%s fayl=%s",
             user_id, sessiya_id, format_, fayl)
    return fayl


def _nakl_word(data: dict, fayl: str) -> None:
    """Nakladnoy Word (.docx) yaratish"""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        raise RuntimeError("python-docx o'rnatilmagan")

    import datetime
    doc = Document()
    doc.add_heading(f"NAKLADNOY № {data.get('nakl_no','—')}", 0)

    doc.add_paragraph(f"Sana: {datetime.datetime.now().strftime('%d.%m.%Y')}")
    doc.add_paragraph(f"Klient: {data['klient'] or 'Noma_lum'}")
    doc.add_paragraph("")

    # Tovarlar jadvali
    headers = ["№", "Nomi", "Miqdor", "Birlik", "Narx", "Jami"]
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        tbl.cell(0, i).text = h
        tbl.cell(0, i).paragraphs[0].runs[0].bold = True

    for idx, t in enumerate(data["tovarlar"], 1):
        row = tbl.add_row()
        row.cells[0].text = str(idx)
        row.cells[1].text = t["nomi"]
        row.cells[2].text = f"{t['miqdor']:.2f}"
        row.cells[3].text = t["birlik"]
        row.cells[4].text = f"{t['narx']:,.0f}"
        row.cells[5].text = f"{t['jami']:,.0f}"

    doc.add_paragraph("")
    doc.add_paragraph(f"Jami: {data['jami_summa']:,.0f} so'm")
    doc.add_paragraph(f"To'langan: {data['tolangan']:,.0f} so'm")
    if data["qarz"] > 0:
        doc.add_paragraph(f"Qarz: {data['qarz']:,.0f} so'm")
    doc.add_paragraph("")

    # Imzo joyi
    sig_tbl = doc.add_table(rows=1, cols=2)
    sig_tbl.style = "Table Grid"
    sig_tbl.cell(0, 0).text = "Beruvchi: _______________"
    sig_tbl.cell(0, 1).text = "Oluvchi:  _______________"
    doc.save(fayl)


def _nakl_excel(data: dict, fayl: str) -> None:
    """Nakladnoy Excel (.xlsx) yaratish"""
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment
    except ImportError:
        raise RuntimeError("openpyxl o'rnatilmagan")

    import datetime
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Nakladnoy"

    ws["A1"] = f"NAKLADNOY № {data.get('nakl_no','—')}"
    ws["A1"].font = Font(bold=True, size=14)
    ws["A2"] = f"Sana: {datetime.datetime.now().strftime('%d.%m.%Y')}"
    ws["A3"] = f"Klient: {data['klient'] or 'Noma_lum'}"

    headers = ["№", "Mahsulot nomi", "Miqdor", "Birlik", "Narx (so'm)", "Jami (so'm)"]
    fill = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=5, column=col, value=h)
        cell.font  = Font(bold=True, color="FFFFFF")
        cell.fill  = fill
        cell.alignment = Alignment(horizontal="center")

    for i, t in enumerate(data["tovarlar"], 1):
        r = 5 + i
        ws.cell(r, 1, i); ws.cell(r, 2, t["nomi"])
        ws.cell(r, 3, t["miqdor"]); ws.cell(r, 4, t["birlik"])
        ws.cell(r, 5, t["narx"]); ws.cell(r, 6, t["jami"])

    last = 5 + len(data["tovarlar"]) + 1
    ws.cell(last, 2, "JAMI:").font = Font(bold=True)
    ws.cell(last, 6, data["jami_summa"]).font = Font(bold=True)
    wb.save(fayl)


def _nakl_pdf(data: dict, fayl: str) -> None:
    """Nakladnoy PDF yaratish (reportlab)"""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
    except ImportError:
        raise RuntimeError("reportlab o'rnatilmagan")

    import datetime
    doc  = SimpleDocTemplate(fayl, pagesize=A4)
    sty  = getSampleStyleSheet()
    els  = []
    els.append(Paragraph(f"NAKLADNOY № {data.get('nakl_no','—')}", sty['Title']))
    els.append(Paragraph(f"Sana: {datetime.datetime.now().strftime('%d.%m.%Y')}", sty['Normal']))
    els.append(Paragraph(f"Klient: {data['klient'] or 'Noma_lum'}", sty['Normal']))
    els.append(Spacer(1, 12))

    tbl_data = [["№", "Mahsulot", "Miqdor", "Narx", "Jami"]]
    for i, t in enumerate(data["tovarlar"], 1):
        tbl_data.append([str(i), t["nomi"],
                          f"{t['miqdor']:.2f} {t['birlik']}",
                          f"{t['narx']:,.0f}", f"{t['jami']:,.0f}"])
    tbl_data.append(["", "JAMI:", "", "", f"{data['jami_summa']:,.0f}"])

    tbl = Table(tbl_data, colWidths=[25, 180, 80, 80, 80])
    tbl.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#1F4E79')),
        ('TEXTCOLOR',  (0,0), (-1,0), colors.white),
        ('FONTNAME',   (0,0), (-1,0), 'Helvetica-Bold'),
        ('GRID',       (0,0), (-1,-1), 0.5, colors.grey),
        ('FONTNAME',   (0,-1),(-1,-1), 'Helvetica-Bold'),
    ]))
    els.append(tbl)
    doc.build(els)




# ════════════════════════════════════════════════════════════════
#  LEDGER RECONCILIATION — Har kuni tekshirish
# ════════════════════════════════════════════════════════════════

@app.task(name="tasks.ledger_reconciliation")
def ledger_reconciliation():
    """Har kuni balans tekshiruvi — DEBIT = CREDIT"""
    import asyncio
    asyncio.get_event_loop().run_until_complete(_ledger_recon_async())


async def _ledger_recon_async():
    from shared.database.pool import pool_init, pool_close, rls_conn
    from shared.services.ledger import balans_tekshir
    await pool_init()
    try:
        users = await _all_faol_users()
        xatolar = []
        for u in users:
            uid = u["id"]
            try:
                async with rls_conn(uid) as c:
                    result = await balans_tekshir(c, uid)
                    if not result["balanslangan"]:
                        xatolar.append(f"⚠️ User {uid}: farq={result['farq']}")
                        log.warning("LEDGER XATO: uid=%d farq=%s", uid, result["farq"])
            except Exception as e:
                log.debug("recon uid=%d: %s", uid, e)

        if xatolar:
            # Admin ga ogohlantirish
            admin_ids = os.environ.get("ADMIN_IDS", "").split(",")
            xabar = (
                "🔴 *LEDGER RECONCILIATION XATO*\n\n"
                + "\n".join(xatolar[:10])
                + f"\n\nJami: {len(xatolar)} ta foydalanuvchida balans buzilgan"
            )
            for aid in admin_ids:
                if aid.strip():
                    try:
                        import httpx
                        bot_token = os.environ.get("BOT_TOKEN", "")
                        if bot_token:
                            async with httpx.AsyncClient(timeout=10) as http:
                                await http.post(
                                    f"https://api.telegram.org/bot{bot_token}/sendMessage",
                                    json={"chat_id": int(aid.strip()), "text": xabar, "parse_mode": "Markdown"}
                                )
                    except Exception:
                        pass
        else:
            log.info("✅ LEDGER RECONCILIATION: barcha userlar balanslangan")
    finally:
        await pool_close()


async def _all_faol_users():
    from shared.database.pool import _pool
    async with _pool().acquire() as c:
        return await c.fetch("SELECT id FROM users WHERE faol=TRUE")
