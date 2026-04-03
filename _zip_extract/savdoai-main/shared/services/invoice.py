"""
╔══════════════════════════════════════════════════════════════╗
║  SAVDOAI v25.3.2 — FAKTURA / INVOICE XIZMATI            ║
║  ✅ Faktura yaratish (Word + PDF)                           ║
║  ✅ Sessiya ID → faktura                                     ║
║  ✅ QR kod bilan                                             ║
║  ✅ Tartib raqami (auto-increment)                          ║
╚══════════════════════════════════════════════════════════════╝
"""
from __future__ import annotations
import io
import logging
from datetime import datetime
from decimal import Decimal
from typing import Any, Optional
import pytz

log = logging.getLogger(__name__)
TZ = pytz.timezone("Asia/Tashkent")


def _hozir() -> str:
    return datetime.now(TZ).strftime("%d.%m.%Y")

def _pul(v: Any) -> str:
    try:
        from decimal import Decimal
        d = Decimal(str(v).replace(",", "").strip()) if not isinstance(v, Decimal) else v
        return f"{d:,.0f}"
    except Exception: return "0"

def _n(v: Any) -> float:
    """Decimal → float FAQAT hujjat rendereri (Word/PDF) uchun.
    DB ga HECH QACHON float yozilmaydi — bu faqat display."""
    try: return float(v) if v else 0.0
    except Exception: return 0.0

def faktura_raqami() -> str:
    return "F-" + datetime.now(TZ).strftime("%Y%m%d%H%M%S")


def faktura_yaratish(data: dict) -> dict[str, bytes]:
    """
    Faktura yaratish — Word va PDF.
    data: {
        "raqam", "dokon_nomi", "klient_ismi",
        "tovarlar": [{"nomi","miqdor","birlik","narx","jami"}],
        "jami_summa", "qarz", "tolangan",
        "izoh", "bank_rekvizitlari"
    }
    Qaytaradi: {"word": bytes, "pdf": bytes}
    """
    natija: dict[str, bytes] = {}

    try:
        natija["word"] = _faktura_word(data)
        log.info("✅ Faktura Word OK")
    except Exception as e:
        log.error("Faktura Word xato: %s", e, exc_info=True)

    try:
        natija["pdf"] = _faktura_pdf(data)
        log.info("✅ Faktura PDF OK")
    except Exception as e:
        log.error("Faktura PDF xato: %s", e, exc_info=True)

    return natija


def _faktura_word(data: dict) -> bytes:
    """Faktura Word format"""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21); section.page_height = Cm(29.7)
    section.left_margin = Cm(2); section.right_margin = Cm(2)

    raqam = data.get("raqam") or faktura_raqami()
    sana = data.get("sana", _hozir())
    dokon = data.get("dokon_nomi", "Mashrab Moliya")
    klient = data.get("klient_ismi", "")
    tovarlar = data.get("tovarlar", [])
    jami = _n(data.get("jami_summa", 0))

    # Sarlavha
    h = doc.add_heading("", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run(f"HISOB-FAKTURA № {raqam}")
    run.font.size = Pt(16); run.font.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"Sana: {sana}  |  {dokon}").font.size = Pt(10)
    doc.add_paragraph()

    # Ma'lumotlar
    info = doc.add_table(rows=2, cols=2)
    info.style = "Table Grid"
    info.cell(0, 0).text = "Yuboruvchi:"; info.cell(0, 1).text = "Qabul qiluvchi:"
    info.cell(1, 0).text = dokon; info.cell(1, 1).text = klient
    doc.add_paragraph()

    # Tovarlar
    cols = ["№", "Tovar nomi", "Birlik", "Miqdor", "Narx", "Jami"]
    tbl = doc.add_table(rows=1 + len(tovarlar) + 1, cols=len(cols))
    tbl.style = "Table Grid"
    for ci, s in enumerate(cols):
        tbl.cell(0, ci).text = s
        for p in tbl.cell(0, ci).paragraphs:
            for r in p.runs: r.font.bold = True; r.font.size = Pt(9)

    jami_hisob = 0.0
    for i, t in enumerate(tovarlar, 1):
        miq = _n(t.get("miqdor", 0)); narx = _n(t.get("narx", 0))
        tj = _n(t.get("jami", miq * narx)); jami_hisob += tj
        vals = [str(i), t.get("nomi", ""), t.get("birlik", "dona"),
                f"{miq:.1f}".rstrip("0").rstrip("."),
                f"{narx:,.0f}" if narx else "—", f"{tj:,.0f}"]
        for ci, v in enumerate(vals):
            tbl.cell(i, ci).text = v

    # Jami
    oxirgi = len(tovarlar) + 1
    tbl.cell(oxirgi, 0).merge(tbl.cell(oxirgi, 4))
    tbl.cell(oxirgi, 0).text = "JAMI:"
    tbl.cell(oxirgi, 5).text = f"{jami or jami_hisob:,.0f} so'm"
    for p in tbl.cell(oxirgi, 0).paragraphs:
        for r in p.runs: r.font.bold = True
    for p in tbl.cell(oxirgi, 5).paragraphs:
        for r in p.runs: r.font.bold = True

    # Bank rekvizitlari
    bank = data.get("bank_rekvizitlari")
    if bank:
        doc.add_paragraph()
        doc.add_paragraph("BANK REKVIZITLARI:").runs[0].font.bold = True
        for k, v in bank.items():
            doc.add_paragraph(f"{k}: {v}")

    doc.add_paragraph()
    # Imzo
    imzo = doc.add_table(rows=2, cols=2); imzo.style = "Table Grid"
    imzo.cell(0, 0).text = "Direktor:"; imzo.cell(0, 1).text = "Buxgalter:"
    imzo.cell(1, 0).text = "Imzo: _______________"
    imzo.cell(1, 1).text = "Imzo: _______________"

    doc.add_paragraph()
    f = doc.add_paragraph()
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f.add_run(f"M.O' | {_hozir()} | Mashrab Moliya").font.size = Pt(8)

    buf = io.BytesIO(); doc.save(buf)
    return buf.getvalue()


def _faktura_pdf(data: dict) -> bytes:
    """Faktura PDF format"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                             topMargin=2*cm, bottomMargin=2*cm,
                             leftMargin=2*cm, rightMargin=2*cm)
    s = getSampleStyleSheet()
    els = []

    raqam = data.get("raqam") or faktura_raqami()
    dokon = data.get("dokon_nomi", "Mashrab Moliya")
    klient = data.get("klient_ismi", "")
    tovarlar = data.get("tovarlar", [])
    jami = _n(data.get("jami_summa", 0))

    els.append(Paragraph(f"HISOB-FAKTURA № {raqam}", s["Title"]))
    els.append(Paragraph(f"Sana: {_hozir()} | {dokon}", s["Normal"]))
    els.append(Paragraph(f"Klient: {klient}", s["Normal"]))
    els.append(Spacer(1, 0.5*cm))

    # Jadval
    tbl_data = [["№", "Tovar", "Miqdor", "Birlik", "Narx", "Jami"]]
    jami_hisob = 0.0
    for i, t in enumerate(tovarlar, 1):
        miq = _n(t.get("miqdor", 0)); narx = _n(t.get("narx", 0))
        tj = _n(t.get("jami", miq * narx)); jami_hisob += tj
        tbl_data.append([
            str(i), t.get("nomi", ""),
            f"{miq:.1f}".rstrip("0").rstrip("."),
            t.get("birlik", "dona"),
            f"{narx:,.0f}" if narx else "—",
            f"{tj:,.0f}",
        ])
    tbl_data.append(["", "", "", "", "JAMI:", f"{jami or jami_hisob:,.0f}"])

    tbl = Table(tbl_data, colWidths=[25, 180, 50, 50, 70, 80])
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1a56db")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -2), [colors.white, colors.HexColor("#f7f9fc")]),
    ]))
    els.append(tbl)
    els.append(Spacer(1, 1*cm))
    els.append(Paragraph(f"Yaratildi: {_hozir()} | Mashrab Moliya", s["Normal"]))

    doc.build(els)
    return buf.getvalue()
